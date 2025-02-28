import binascii
import io
import json

import docx
import np

misperrors = {"error": "Error"}
mispattributes = {"input": ["attachment"], "output": ["freetext", "text"]}
moduleinfo = {
    "version": "0.1",
    "author": "Sascha Rommelfangen",
    "description": "Module to extract freetext from a .docx document.",
    "module-type": ["expansion"],
    "name": "DOCX Enrich",
    "logo": "docx.png",
    "requirements": ["docx python library"],
    "features": (
        "The module reads the text contained in a .docx document. The result is passed to the freetext import parser so"
        " IoCs can be extracted out of it."
    ),
    "references": [],
    "input": "Attachment attribute containing a .docx document.",
    "output": "Text and freetext parsed from the document.",
}

moduleconfig = []


def handler(q=False):
    if q is False:
        return False
    q = json.loads(q)
    filename = q["attachment"]
    try:
        docx_array = np.frombuffer(binascii.a2b_base64(q["data"]), np.uint8)
    except Exception as e:
        print(e)
        err = "Couldn't fetch attachment (JSON 'data' is empty). Are you using the 'Query enrichment' action?"
        misperrors["error"] = err
        print(err)
        return misperrors

    doc_content = ""
    doc_file = io.BytesIO(docx_array)
    try:
        doc = docx.Document(doc_file)
        for para in doc.paragraphs:
            print(para.text)
            doc_content = doc_content + "\n" + para.text
        tables = doc.tables
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        print(para.text)
                        doc_content = doc_content + "\n" + para.text
        print(doc_content)
        return {
            "results": [
                {
                    "types": ["freetext"],
                    "values": doc_content,
                    "comment": ".docx-to-text from file " + filename,
                },
                {
                    "types": ["text"],
                    "values": doc_content,
                    "comment": ".docx-to-text from file " + filename,
                },
            ]
        }
    except Exception as e:
        print(e)
        err = "Couldn't analyze file as .docx. Error was: " + str(e)
        misperrors["error"] = err
        return misperrors


def introspection():
    return mispattributes


def version():
    moduleinfo["config"] = moduleconfig
    return moduleinfo
