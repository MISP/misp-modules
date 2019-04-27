import json
import binascii
import np
import docx
import io

misperrors = {'error': 'Error'}
mispattributes = {'input': ['attachment'],
                  'output': ['freetext', 'text']}
moduleinfo = {'version': '0.1', 'author': 'Sascha Rommelfangen',
              'description': '.docx to freetext-import IOC extractor',
              'module-type': ['expansion']}

moduleconfig = []


def handler(q=False):
    if q is False:
        return False
    q = json.loads(q)
    filename = q['attachment']
    try:
        docx_array = np.frombuffer(binascii.a2b_base64(q['data']), np.uint8)
    except Exception as e:
        print(e)
        err = "Couldn't fetch attachment (JSON 'data' is empty). Are you using the 'Query enrichment' action?"
        misperrors['error'] = err
        print(err)
        return misperrors

    doc_content = ""
    doc_file = io.BytesIO(docx_array)
    try:
        doc = docx.Document(doc_file)
        for para in doc.paragraphs:
            print(para.text)
            doc_content = doc_content + "\n" + para.text
        tables = doc.tables
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        print(para.text)
                        doc_content = doc_content + "\n" + para.text
        print(doc_content)
        return {'results': [{'types': ['freetext'], 'values': doc_content, 'comment': ".docx-to-text from file " + filename},
                            {'types': ['text'], 'values': doc_content, 'comment': ".docx-to-text from file " + filename}]}
    except Exception as e:
        print(e)
        err = "Couldn't analyze file as .docx. Error was: " + str(e)
        misperrors['error'] = err
        return misperrors


def introspection():
    return mispattributes


def version():
    moduleinfo['config'] = moduleconfig
    return moduleinfo
